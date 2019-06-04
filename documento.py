from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from docx.shared import Pt, Inches
from docx.dml.color import ColorFormat
from docx.shared import RGBColor, Cm
from auxiliares import devolverMes
from estilo import Estilos
from elastic import Info
from grafico import Grafico


class Word:

    def abrirDocumento(self):
        informe = Document('./template-info/template.docx')
        estilos = Estilos()
        estilos.styleTitle(informe)
        estilos.styleBodyText(informe)
        estilos.styleHeading1(informe)
        estilos.styleSubtitle(informe)
        return informe


    def guarda(self, informe):
        informe.save('./informe/informe.docx')

    def titulo(self, informe, title, subtitle):

        informe.add_paragraph(title, style = 'Title')
        informe.add_paragraph(subtitle, style = 'Subtitle')
        informe.add_page_break()

    def creaPagina(self, informe, response, title, introduction, column_tags):

        informe.add_paragraph(title, style = 'Heading 1')
        informe.add_paragraph(introduction, style='Body Text')
        self.tabla(informe, column_tags, response)

        grafico = Grafico().crearGrafico([i.doc_count for i in response], [i.key for i in response], filename='graficoDeBarra.png')
        informe.add_picture('./imagenes/graficoDeBarra.png')
        informe.add_page_break()

    def tabla(self, informe, column_tags, informacion):
        table = informe.add_table(rows = 1, cols = len(column_tags))
        table.style = "BTR Table"
        table.autofit = False
        table.width = Cm(4.19)
        column_etiquetas = table.rows[0].cells
        index = 0
        for cell in table.columns[0].cells:
            cell.width = Cm(4.19)

        for cell in table.columns[1].cells:
            cell.width = Cm(4.19)

        for i in column_etiquetas:
            i.text = column_tags[index]
            index += 1

        diccionario_informacion = dict(zip([i.key for i in informacion],[i.doc_count for i in informacion]))



        for key, value in diccionario_informacion.items():
            row = table.add_row()
            cells = row.cells
            cells[0].text = key.capitalize()
            cells[1].text = str(value)
